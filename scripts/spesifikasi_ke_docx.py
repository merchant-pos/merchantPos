"""Mengubah SPESIFIKASI-KAATAGO.md jadi dokumen Word resmi.

Sengaja dibuat sebagai konverter, bukan penulis dokumen kedua: isi yang
ditulis dua kali akan berbeda pada perubahan berikutnya, dan yang paling
mungkin ketinggalan justru berkas Word-nya — yang justru itu yang
dibagikan ke penguji.

Yang ditangani hanya sebagian Markdown yang benar-benar dipakai berkas
itu: judul, paragraf, daftar, tabel pipa, kutipan, blok kode, gambar,
serta penegasan tebal/miring/kode di dalam baris.
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Jalur dihitung dari letak berkas ini, bukan ditulis mutlak — supaya
# repo yang di-clone ke folder lain tetap bisa membangun dokumennya.
#
#   /usr/bin/python3 scripts/spesifikasi_ke_docx.py [nama-berkas-md]
#
# Tanpa argumen: membangun SPESIFIKASI-KAATAGO.md. Judul sampul dan
# keterangan versinya dibaca dari berkas Markdown-nya sendiri, jadi satu
# skrip ini melayani semua dokumen tanpa perlu disunting tiap kali ada
# dokumen baru.
#
# Butuh: python3 -m pip install --user python-docx
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOC = sys.argv[1] if len(sys.argv) > 1 else "SPESIFIKASI-KAATAGO.md"
if not DOC.endswith(".md"):
    DOC += ".md"
SRC = os.path.join(REPO, "docs", DOC)
OUT = SRC[:-3] + ".docx"
IMG_ROOT = os.path.dirname(SRC)

BRAND = RGBColor(0x4F, 0x46, 0xE5)
BRAND_DARK = RGBColor(0x37, 0x30, 0xA3)
INK = RGBColor(0x14, 0x15, 0x2A)
MUTED = RGBColor(0x6E, 0x72, 0x8C)

HEADER_SHADE = "4F46E5"
ZEBRA_SHADE = "F4F5FB"
QUOTE_SHADE = "FFF7E6"
CODE_SHADE = "F2F3F8"


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def para_shade(p, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    p._p.get_or_add_pPr().append(el)


def left_bar(p, hexcolor):
    """Garis tebal di tepi kiri — penanda kutipan/peringatan."""
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hexcolor)
    pbdr.append(left)
    p._p.get_or_add_pPr().append(pbdr)


def field(paragraph, instr, placeholder="—"):
    """Menyisipkan field Word (dipakai untuk nomor halaman dan daftar isi).

    [placeholder] adalah yang terlihat sampai Word memperbarui fieldnya.
    Word menghitung ulang PAGE/NUMPAGES sendiri saat dokumen dibuka,
    tapi TOC menunggu diminta — jadi placeholder-nya perlu berupa
    kalimat yang menjelaskan itu, bukan tanda hubung yang membingungkan.
    """
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)

    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._r.append(it)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)

    r = paragraph.add_run(placeholder)
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = MUTED

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, base_size=10.5, color=None):
    """Menuliskan teks berikut penegasan tebal/miring/kode di dalamnya."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = BRAND_DARK
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("["):
            label = part[1 : part.index("]")]
            r = paragraph.add_run(label)
            r.font.color.rgb = BRAND
        else:
            r = paragraph.add_run(part)
        r.font.size = Pt(base_size)
        if color is not None and r.font.color.rgb is None:
            r.font.color.rgb = color


def page_break(doc):
    """Pemisah halaman selalu di paragraf sendiri.

    Menempelkannya ke paragraf terakhir tidak bisa diandalkan: tabel
    bukan paragraf, jadi setelah sebuah tabel `doc.paragraphs[-1]` masih
    menunjuk teks di atasnya — dan halamannya terputus di tempat yang
    salah.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def read_front_matter(path):
    """Judul, subjudul, dan keterangan versi dari kepala Markdown-nya.

    Dibaca dari berkasnya, bukan ditulis di skrip ini: dokumen kedua yang
    memaksa skripnya disunting adalah dokumen yang sampulnya akan salah
    pada perubahan berikutnya.
    """
    title, subtitle, meta = "MerchantPOS", "", []
    for line in open(path, encoding="utf-8").read().split("\n"):
        st = line.strip()
        if st.startswith("# ") and title == "MerchantPOS":
            heading = st[2:].strip()
            if "—" in heading:
                title, subtitle = [x.strip() for x in heading.split("—", 1)]
            else:
                title = heading
        m = re.match(r"^\*\*(.+?):\*\*\s*(.+)$", st)
        if m:
            meta.append((m.group(1), m.group(2)))
        if st.startswith("## "):
            break
    return title, subtitle, meta


def build():
    doc = Document()
    doc_title, doc_subtitle, doc_meta = read_front_matter(SRC)

    # ── Gaya dasar ────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 18, BRAND_DARK),
        ("Heading 2", 14, BRAND_DARK),
        ("Heading 3", 12, INK),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    # ── Sampul ────────────────────────────────────────────────────────
    logo = os.path.join(IMG_ROOT, "gambar", "logo-merchantpos.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    if os.path.exists(logo):
        p.add_run().add_picture(logo, width=Cm(4.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MerchantPOS")
    r.font.size = Pt(34)
    r.bold = True
    r.font.color.rgb = BRAND_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_subtitle.upper() or "SPECIFICATION DOCUMENT")
    r.font.size = Pt(15)
    r.bold = True
    r.font.color.rgb = INK

    # Baris kecil di bawah judulnya hanya muncul kalau memang menambah
    # keterangan. Mengulang nama mereknya persis di bawah nama merek
    # yang sudah tercetak besar cuma jadi gema.
    if doc_title and doc_title.lower() != "merchantpos":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(doc_title)
        r.font.size = Pt(11.5)
        r.font.color.rgb = MUTED

    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    # Kendali dokumen
    ctrl = [("Nama Aplikasi", "MerchantPOS — Kasir & Self-Order untuk Resto")]
    ctrl += doc_meta
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in ctrl:
        row = t.add_row()
        row.cells[0].width = Cm(4.6)
        row.cells[1].width = Cm(10.4)
        rp = row.cells[0].paragraphs[0]
        rr = rp.add_run(k)
        rr.bold = True
        rr.font.size = Pt(10)
        shade(row.cells[0], ZEBRA_SHADE)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(10)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)

    page_break(doc)

    # ── Kepala & kaki halaman ─────────────────────────────────────────
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version = next((v for k, v in doc_meta if "Versi Aplikasi" in k), "")
    r = hp.add_run(
        "MerchantPOS — ${}{}".format(
            doc_subtitle or "Specification Document",
            "  |  v" + version.split(" ")[0] if version else "",
        ).replace("$", "")
    )
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Halaman ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    field(fp, "PAGE")
    r = fp.add_run(" dari ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    field(fp, "NUMPAGES")
    for r in fp.runs:
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED

    # ── Daftar isi otomatis ───────────────────────────────────────────
    h = doc.add_heading("Daftar Isi", level=1)
    p = doc.add_paragraph()
    field(
        p,
        r'TOC \o "1-3" \h \z \u',
        placeholder="Daftar isi dimuat setelah field diperbarui.",
    )
    note = doc.add_paragraph()
    r = note.add_run(
        "Klik kanan pada daftar di atas → Update Field, untuk memuat nomor halamannya."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    page_break(doc)

    # ── Isi, dari Markdown ────────────────────────────────────────────
    lines = open(SRC, encoding="utf-8").read().split("\n")
    i = 0
    skip_toc = False
    seen_section = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Judul dokumen dan daftar isi manual dilewati: sampul dan field
        # TOC sudah menggantikan keduanya.
        if stripped.startswith("# "):
            i += 1
            continue
        # Keterangan versi/tanggal/status sudah tercetak di tabel kendali
        # dokumen pada sampulnya; mengulangnya di badan dokumen membuat
        # dua tempat yang harus diperbarui bersamaan.
        if re.match(r"^\*\*[^*]+:\*\*\s", stripped) and not seen_section:
            i += 1
            continue
        if stripped == "## Daftar Isi":
            skip_toc = True
            i += 1
            continue
        if skip_toc:
            if stripped.startswith("## "):
                skip_toc = False
            else:
                i += 1
                continue

        if not stripped or stripped == "---":
            i += 1
            continue

        # Galeri tangkapan layar: baris "!!ss[caption](path)" yang
        # berurutan dikumpulkan jadi satu tabel tanpa garis, tiga kolom.
        #
        # Ditaruh berkelompok, bukan satu per halaman: dua ratus tangkapan
        # layar seukuran halaman penuh menghasilkan dokumen yang tidak
        # akan pernah dibuka sampai habis. Bertiga sebaris, ukurannya
        # masih cukup untuk mengenali layar mana yang sedang dilihat, dan
        # seluruh peran muat dibaca dalam beberapa halaman.
        if stripped.startswith("!!ss["):
            grup = []
            while i < len(lines) and lines[i].strip().startswith("!!ss["):
                mm = re.match(r"!!ss\[(.*?)\]\((.+?)\)", lines[i].strip())
                if mm:
                    path = os.path.normpath(os.path.join(IMG_ROOT, mm.group(2)))
                    if os.path.exists(path):
                        grup.append((mm.group(1), path))
                    else:
                        # Diam-diam melewatkannya berarti lampiran yang
                        # kurang setengahnya tanpa ada yang tahu — dan
                        # itu memang sempat terjadi: nama folder yang
                        # mengandung tanda kurung memutus jalurnya di
                        # tengah.
                        print("  ! gambar tidak ditemukan:", path)
                i += 1
            if grup:
                add_gallery(doc, grup)
            continue

        # Gambar
        m = re.match(r"!\[(.*?)\]\((.+?)\)", stripped)
        if m:
            caption, rel = m.group(1), m.group(2)
            path = os.path.join(IMG_ROOT, rel)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                width = Cm(11) if "logo" in rel or "qr-meja" in rel else Cm(15.5)
                p.add_run().add_picture(path, width=width)
                if caption:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(caption)
                    cr.italic = True
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = MUTED
            i += 1
            continue

        # Judul bagian
        # "##" jadi Heading 1: di Markdown tingkat teratas dipakai judul
        # dokumen, sedangkan di Word judulnya sudah ada di sampul — jadi
        # bagian bernomor 1..13 itulah tingkat teratas yang sebenarnya.
        if stripped.startswith("### "):
            doc.add_heading(clean(stripped[4:]), level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            seen_section = True
            doc.add_heading(clean(stripped[3:]), level=1)
            i += 1
            continue

        # Blok kode
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            para_shade(p, CODE_SHADE)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = BRAND_DARK
            continue

        # Kutipan / peringatan
        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            para_shade(p, QUOTE_SHADE)
            left_bar(p, "F59E0B")
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(6)
            add_inline(p, " ".join(buf), base_size=10)
            continue

        # Tabel
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            add_table(doc, block)
            continue

        # Daftar berbutir / bernomor
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent, marker, text = m.group(1), m.group(2), m.group(3)
            i += 1
            # Baris lanjutan dari butir yang sama
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip() and not re.match(r"^(\s*)([-*]|\d+\.)\s+", lines[i]) and not lines[i].strip().startswith("|"):
                text += " " + lines[i].strip()
                i += 1
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.8 + len(indent) * 0.2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            continue

        # Paragraf biasa — baris berturut digabung
        buf = [stripped]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].strip().startswith(("#", "|", ">", "-", "*", "```", "!["))
            and not re.match(r"^\d+\.\s", lines[i].strip())
        ):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(buf))

    doc.save(OUT)
    print("ditulis:", OUT)


def clean(text):
    return re.sub(r"[*`]", "", text).strip()


def add_gallery(doc, items, per_row=3):
    """Menyusun tangkapan layar jadi tabel tanpa garis.

    Tabel, bukan deretan paragraf: paragraf gambar berukuran sama akan
    dijatuhkan satu per baris oleh Word, dan dua ratus gambar berarti
    dua ratus baris. Sel tabel menahan tiganya tetap sebaris.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    lebar = Cm(5.0)
    for mulai in range(0, len(items), per_row):
        potong = items[mulai:mulai + per_row]
        t = doc.add_table(rows=2, cols=per_row)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for kolom in range(per_row):
            sel_gambar = t.cell(0, kolom)
            sel_teks = t.cell(1, kolom)
            pg = sel_gambar.paragraphs[0]
            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pg.paragraph_format.space_after = Pt(0)
            pt = sel_teks.paragraphs[0]
            pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt.paragraph_format.space_after = Pt(10)
            if kolom < len(potong):
                caption, path = potong[kolom]
                try:
                    pg.add_run().add_picture(path, width=lebar)
                except Exception:
                    # Berkas rusak tidak boleh menjatuhkan seluruh
                    # dokumen — lampiran yang kurang satu gambar masih
                    # jauh lebih berguna daripada dokumen yang gagal
                    # dibuat.
                    continue
                r = pt.add_run(caption)
                r.italic = True
                r.font.size = Pt(7.5)
                r.font.color.rgb = MUTED


def add_table(doc, block):
    rows = []
    for ln in block:
        if re.match(r"^\|[\s:\-|]+\|$", ln):  # baris pemisah
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.autofit = True

    for ri, cells in enumerate(rows):
        row = t.add_row()
        for ci in range(ncol):
            cell = row.cells[ci]
            text = cells[ci] if ci < len(cells) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if ri == 0:
                shade(cell, HEADER_SHADE)
                r = p.add_run(clean(text))
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # Baris berselang-seling: tabel matriks di dokumen ini
                # lebarnya sampai delapan kolom, dan tanpa selang-seling
                # mata gampang meleset satu baris saat membacanya.
                if ri % 2 == 0:
                    shade(cell, ZEBRA_SHADE)
                add_inline(p, text, base_size=9.5)
                if text.strip() in {"✔", "–", "-", "✘"}:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ulangi baris judul di tiap halaman — tabel panjang yang terpotong
    # halaman tanpa judul praktis tidak terbaca.
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


if __name__ == "__main__":
    build()
